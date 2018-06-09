#! python3
# -- coding: utf-8

from docx import Document
from docx.shared import Inches
import logging
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import os, sys
from docx.enum.text import WD_ALIGN_PARAGRAPH


# 这个类不保存文本数据，只接受文本数据，通过给定的格式信息和文本的定位直接生成和写入到docx中去
class  CreTranslationDocx :
    def __init__(self,fileDir,fileName):
            self.fileDir = fileDir
            self.fileName = fileName
            self.totalFilePath = fileDir+"/"+fileName
            # 不检查目录 fileDir 是否存在，只检查 文件  fileName是否存在
            if os.path.exists(self.totalFilePath) == True:
                os.remove(self.totalFilePath)
            self.document = Document( )
    # 默认设置3级标题，和文本字体一样大小，居中
    def add_heading(self,text_for_add):
            # self.document.add_heading(text_for_add,3)
            # paragraph = document.add_paragraph(text_for_add)
            paragraph = self.document.add_paragraph('') 
            run = paragraph.add_run(text_for_add)
            run.bold =  True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 文章内部小标题
    def add_title(self,text_for_add):
            paragraph = self.document.add_paragraph('\t') 
            run = paragraph.add_run(text_for_add)
            run.font.size = Pt(12)
            run.bold =  True

    def add_paragraph(self):
            paragraph = self.document.add_paragraph('\t') 
            return paragraph

    def add_sentence(self,text_for_add,paragraph):
            run = paragraph.add_run(text_for_add)
            #设置字号
            run.font.size = Pt(12)
            #设置中文字体
            run.font.name='宋体'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            #设置粗体
            # run = paragraph.add_run('粗体').bold = True          


    def save_(self):
            #保存文件
            self.document.save(self.totalFilePath)
            # document




def testDocx1():
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )

    document.add_picture('monty-truth.png', width=Inches(1.25))

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    # for item in recordset:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(item.qty)
    #     row_cells[1].text = str(item.id)
    #     row_cells[2].text = item.desc

    document.add_page_break()

    document.save('demo.docx')

def testDocx2():
    #打开文档
    document = Document()
    #加入不同等级的标题
    document.add_heading('MS WORD写入测试',0)
    document.add_heading('一级标题',1)
    document.add_heading('二级标题',2)

    # 3级的字体号就和正文一样了
    document.add_heading(u'3级标题',3)
    #添加文本

    paragraph = document.add_paragraph(u'我们在做文本测试！')
 

    #设置字号
    run = paragraph.add_run(u'设置字号、')
    run.font.size = Pt(24)
    #设置字体
    run = paragraph.add_run('Set Font,')
    run.font.name = 'Consolas'

    #设置中文字体
    run = paragraph.add_run(u'设置中文字体、')
    run.font.name=u'宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    #设置斜体
    run = paragraph.add_run(u'斜体、')
    run.italic = True

    #设置粗体
    run = paragraph.add_run(u'粗体').bold = True

    #增加引用
    document.add_paragraph('Intense quote', style='Intense Quote')

    #增加无序列表
    document.add_paragraph(
        u'无序列表元素1', style='List Bullet'
    )
    document.add_paragraph(
        u'无序列表元素2', style='List Bullet'
    )
    #增加有序列表
    document.add_paragraph(
        u'有序列表元素1', style='List Number'
    )
    document.add_paragraph(
        u'有序列表元素2', style='List Number'
    )
    #增加图像（此处用到图像image.bmp，请自行添加脚本所在目录中）
    # document.add_picture('image.bmp', width=Inches(1.25))

    #增加表格
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    #再增加3行表格元素
    # for i in xrange(3):
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = 'test'+str(i)
    #     row_cells[1].text = str(i)
    #     row_cells[2].text = 'desc'+str(i)

    #增加分页
    document.add_page_break()
    document.add_page_break()

    #保存文件
    document.save(u'测试.docx')


def main():      

    logging.basicConfig(filename='example.log', level=logging.DEBUG)

    # testDocx2()
    creDoc = CreTranslationDocx("/home/jty/codetest","test.docx")
    creDoc.add_heading( "标题")
   
    # creDoc.add_sentence("话的意思就是，当模块被直接运行时，以下“ss"代码块将被运行", creDoc.add_paragraph())
    # creDoc.add_sentence("话的意思就是，当模块被直接运行时，以下代码块将被运行", creDoc.add_paragraph())
    creDoc.add_title("这句话的意思就是")
    creDoc.add_sentence("话的意思就是，当模块被直接运行时，以下代码块将被运行", creDoc.add_paragraph())
    creDoc.save_()


# __name__ 是当前模块名，当模块被直接运行时模块名为 __main__ 。这句话的意思就是，当模块被直接运行时，以下代码块将被运行，当模块是被导入时，代码块不被运行。
if __name__ == "__main__":      
    main()   























